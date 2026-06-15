import base64
import io
from pathlib import Path

import fitz
from docx import Document
from PIL import Image

SUPPORTED_IMAGE = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif"}
SUPPORTED_PDF = {".pdf"}
SUPPORTED_DOC = {".docx", ".doc"}


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in SUPPORTED_IMAGE:
        return "image"
    if ext in SUPPORTED_PDF:
        return "pdf"
    if ext in SUPPORTED_DOC:
        return "doc"
    return "unknown"


def extract_text_from_pdf(data: bytes) -> str:
    text_parts = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()


def pdf_to_images(data: bytes, max_pages: int = 5) -> list[dict]:
    images = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            images.append({
                "mime": "image/png",
                "data": base64.b64encode(img_data).decode("utf-8"),
            })
    return images


def image_to_payload(data: bytes) -> dict:
    img = Image.open(io.BytesIO(data))
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return {
        "mime": "image/jpeg",
        "data": base64.b64encode(buf.getvalue()).decode("utf-8"),
    }


def extract_content(filename: str, data: bytes) -> dict:
    file_type = get_file_type(filename)
    result = {"file_type": file_type, "text": "", "images": []}

    if file_type == "pdf":
        result["text"] = extract_text_from_pdf(data)
        result["images"] = pdf_to_images(data)
    elif file_type == "doc":
        result["text"] = extract_text_from_docx(data)
    elif file_type == "image":
        result["images"] = [image_to_payload(data)]
    else:
        raise ValueError(f"不支援的檔案格式: {Path(filename).suffix}")

    return result